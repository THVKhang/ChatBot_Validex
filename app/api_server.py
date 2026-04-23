from collections import defaultdict
from collections import deque
import asyncio
import importlib
import json
import logging
import time
from typing import Any, Literal
from uuid import uuid4

from fastapi import FastAPI
from fastapi import Header
from fastapi import HTTPException
from fastapi import Request
from fastapi import Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from pydantic import Field

from app.config import settings
from app.langchain_pipeline import pipeline
from app.main import process_prompt
from app.prompt_guard import validate_prompt
from app.publisher import build_publish_output
from app.report_store import delete_report
from app.report_store import get_report
from app.report_store import list_reports
from app.report_store import save_report
from app.report_store import update_report_status
from app.session_manager import SessionManager
from app.session_store import load_session
from app.session_store import save_session as persist_session
from app.source_analytics import fetch_knowledge_health
from app.source_analytics import fetch_source_analytics

logger = logging.getLogger(__name__)


class ChatRequest(BaseModel):
    prompt: str
    session_id: str | None = None


class BlogSectionPayload(BaseModel):
    heading: str
    body: str
    image_url: str
    image_alt: str


class GeneratedPayload(BaseModel):
    title: str
    outline: list[str] = Field(default_factory=list)
    draft: str
    sources_used: list[str] = Field(default_factory=list)
    sections: list[BlogSectionPayload] = Field(default_factory=list)


class SaveReportRequest(BaseModel):
    prompt: str
    generated: GeneratedPayload
    session_id: str | None = None


class UpdateReportStatusRequest(BaseModel):
    status: Literal["Draft", "Reviewed", "Approved"]


app = FastAPI(title="AI Blog Generator API", version="0.1.0")

_allowed_origins = [o.strip() for o in settings.allowed_origins.split(",") if o.strip()]
app.add_middleware(
    CORSMiddleware,
    allow_origins=_allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

try:
    from app.auth import auth_router, get_current_user_id
    app.include_router(auth_router)
except ImportError:
    # Handle tests/mock cases
    def get_current_user_id(): return None

# TTL-tracked sessions: maps session_id -> (SessionManager, last_access_timestamp)
_sessions_store: dict[str, tuple[SessionManager, float]] = {}
_rate_limit_window: dict[str, deque[float]] = defaultdict(deque)


def _get_or_create_session(session_id: str, user_id: int | None = None) -> SessionManager:
    """Fetch an existing session or create a new one, updating last-access time.

    Tries in-memory cache first, then PostgreSQL, then creates a new session.
    """
    ttl = max(60, settings.session_ttl_seconds)
    now = time.time()
    # Evict stale sessions to prevent memory growth.
    stale = [sid for sid, (_, ts) in _sessions_store.items() if now - ts > ttl]
    for sid in stale:
        _sessions_store.pop(sid, None)
    if session_id in _sessions_store:
        mgr, _ = _sessions_store[session_id]
        _sessions_store[session_id] = (mgr, now)
        return mgr
    # Try loading from PostgreSQL
    mgr = load_session(session_id, user_id=user_id)
    if mgr is None:
        mgr = SessionManager()
    _sessions_store[session_id] = (mgr, now)
    return mgr


async def _save_session_async(session_id: str, session: SessionManager, user_id: int | None = None) -> None:
    """Persist session to DB in a background thread (fire-and-forget)."""
    try:
        await asyncio.to_thread(persist_session, session_id, session, user_id)
    except Exception:
        pass  # Non-critical — in-memory session still works

_metrics: dict[str, object] = {
    "chat_requests_total": 0,
    "chat_errors_total": 0,
    "quality_gate_blocked_total": 0,
    "generation_mode_count": defaultdict(int),
    "retrieval_mode_count": defaultdict(int),
    "latency_ms_samples": deque(maxlen=max(20, settings.metrics_window_size)),
}


def _redis_client():
    if not settings.use_redis_rate_limit:
        return None
    try:
        redis_module = importlib.import_module("redis")
    except Exception:
        return None
    redis_cls = getattr(redis_module, "Redis", None)
    if redis_cls is None:
        return None
    try:
        return redis_cls.from_url(settings.redis_url)
    except Exception:
        return None


_redis = _redis_client()


@app.middleware("http")
async def rate_limit_middleware(request: Request, call_next):
    if not settings.use_rate_limit:
        return await call_next(request)

    # Apply rate limit to both chat endpoints
    if request.url.path not in ("/api/chat", "/api/chat/stream"):
        return await call_next(request)

    client_id = request.client.host if request.client else "unknown"
    if _redis is not None:
        key = f"rate_limit:{client_id}:{int(time.time() // 60)}"
        try:
            count = _redis.incr(key)
            if count == 1:
                _redis.expire(key, 61)
            if count > max(1, settings.rate_limit_per_minute):
                return JSONResponse(
                    status_code=429,
                    content={"detail": "Rate limit exceeded. Please retry in a minute."},
                )
        except Exception:
            pass

    now = time.time()
    window = _rate_limit_window[client_id]
    while window and now - window[0] > 60:
        window.popleft()

    if len(window) >= max(1, settings.rate_limit_per_minute):
        return JSONResponse(
            status_code=429,
            content={"detail": "Rate limit exceeded. Please retry in a minute."},
        )

    window.append(now)
    return await call_next(request)


@app.middleware("http")
async def request_id_middleware(request: Request, call_next):
    """Attach a unique request_id to each request for tracing."""
    request_id = request.headers.get("X-Request-ID") or str(uuid4())
    request.state.request_id = request_id
    response = await call_next(request)
    response.headers["X-Request-ID"] = request_id
    return response


def _validate_chat_prompt(prompt: str) -> str:
    """Run prompt guard and raise 400 if invalid. Returns cleaned prompt."""
    result = validate_prompt(prompt)
    if not result.is_valid:
        raise HTTPException(status_code=400, detail=result.rejection_reason or "Invalid prompt.")
    if result.warnings:
        logger.info("prompt_guard.warnings: %s", result.warnings)
    return result.cleaned_prompt


@app.get("/api/health")
def health() -> dict:
    runtime = pipeline.runtime_status()
    return {
        "status": "ok",
        "runtime": runtime,
        "flags": {
            "use_live_llm": settings.use_live_llm,
            "use_pinecone_retrieval": settings.use_pinecone_retrieval,
            "use_agentic_rag": settings.use_agentic_rag,
            "use_rate_limit": settings.use_rate_limit,
            "use_redis_rate_limit": settings.use_redis_rate_limit,
        },
    }


@app.get("/api/metrics")
def metrics() -> dict:
    latency_samples = list(_metrics["latency_ms_samples"])
    avg_latency = round(sum(latency_samples) / len(latency_samples), 2) if latency_samples else 0.0
    p95_latency = 0.0
    if latency_samples:
        ordered = sorted(latency_samples)
        index = int(0.95 * (len(ordered) - 1))
        p95_latency = round(float(ordered[index]), 2)

    return {
        "chat_requests_total": _metrics["chat_requests_total"],
        "chat_errors_total": _metrics["chat_errors_total"],
        "quality_gate_blocked_total": _metrics["quality_gate_blocked_total"],
        "generation_mode_count": dict(_metrics["generation_mode_count"]),
        "retrieval_mode_count": dict(_metrics["retrieval_mode_count"]),
        "latency": {
            "samples": len(latency_samples),
            "avg_ms": avg_latency,
            "p95_ms": p95_latency,
        },
    }


@app.get("/api/chat/sessions")
def get_chat_sessions(limit: int = 50, user_id: int | None = Depends(get_current_user_id)) -> list[dict]:
    """Retrieve a list of recent chat sessions for the sidebar."""
    from app.session_store import list_sessions
    return list_sessions(limit, user_id=user_id)


@app.get("/api/chat/sessions/{session_id}")
def get_chat_session(session_id: str, user_id: int | None = Depends(get_current_user_id)) -> dict:
    """Retrieve full history for a specific session."""
    from app.session_store import load_session
    session = load_session(session_id, user_id=user_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found or unauthorized")
    
    turns = []
    for turn in session.turns:
        turns.append({
            "user_prompt": turn.user_prompt,
            "assistant_output": turn.assistant_output,
            "parsed_intent": turn.parsed_intent,
            "parsed_topic": turn.parsed_topic,
        })
    return {"session_id": session_id, "turns": turns}


class ExportRequest(BaseModel):
    markdown: str
    format: str = "docx"  # "docx" or "html"

from fastapi.responses import Response

@app.post("/api/chat/export")
def export_chat(request: ExportRequest) -> Response:
    """Export markdown content to docx or html."""
    if request.format == "docx":
        try:
            from docx import Document
            doc = Document()
            for line in request.markdown.split('\n'):
                line_stripped = line.strip()
                if line_stripped.startswith('# '):
                    doc.add_heading(line_stripped[2:].strip(), 1)
                elif line_stripped.startswith('## '):
                    doc.add_heading(line_stripped[3:].strip(), 2)
                elif line_stripped.startswith('### '):
                    doc.add_heading(line_stripped[4:].strip(), 3)
                elif line_stripped:
                    doc.add_paragraph(line_stripped)
            
            import io
            f = io.BytesIO()
            doc.save(f)
            return Response(
                content=f.getvalue(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=validex_report.docx"}
            )
        except Exception as exc:
            logger.error("Docx export failed: %s", exc)
            raise HTTPException(status_code=500, detail="Docx export failed")
    elif request.format == "html":
        try:
            import markdown2
            html = markdown2.markdown(request.markdown)
            return Response(
                content=html,
                media_type="text/html",
                headers={"Content-Disposition": "attachment; filename=validex_report.html"}
            )
        except Exception as exc:
            logger.error("HTML export failed: %s", exc)
            raise HTTPException(status_code=500, detail="HTML export failed")
    else:
        raise HTTPException(status_code=400, detail="Invalid format")


from fastapi import UploadFile, File

@app.post("/api/chat/upload")
async def chat_upload(file: UploadFile = File(...)) -> dict:
    """Extract text from uploaded PDF/Docx to serve as context."""
    filename = file.filename or "unknown"
    content_bytes = await file.read()
    extracted_text = ""
    
    if filename.lower().endswith(".pdf"):
        try:
            import fitz
            doc = fitz.open(stream=content_bytes, filetype="pdf")
            for page in doc:
                extracted_text += page.get_text() + "\n"
        except Exception as exc:
            logger.error("PDF extraction failed: %s", exc)
            raise HTTPException(status_code=400, detail="Failed to extract PDF text")
    elif filename.lower().endswith(".docx"):
        try:
            import io
            from docx import Document
            doc = Document(io.BytesIO(content_bytes))
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
        except Exception as exc:
            logger.error("DOCX extraction failed: %s", exc)
            raise HTTPException(status_code=400, detail="Failed to extract DOCX text")
    else:
        extracted_text = content_bytes.decode("utf-8", errors="ignore")
        
    # We return the extracted text to the frontend so it can be appended to the prompt.
    return {
        "filename": filename,
        "extracted_text": extracted_text.strip()
    }



@app.post("/api/chat")
async def chat(request: ChatRequest, req: Request = None, user_id: int | None = Depends(get_current_user_id)) -> dict:
    cleaned_prompt = _validate_chat_prompt(request.prompt)
    request_id = getattr(req.state, "request_id", None) if req else None
    start = time.perf_counter()
    session_id = request.session_id or str(uuid4())
    session = _get_or_create_session(session_id, user_id=user_id)
    _metrics["chat_requests_total"] += 1

    try:
        payload = await asyncio.to_thread(
            process_prompt, cleaned_prompt, session, request_id=request_id,
        )
    except Exception:
        _metrics["chat_errors_total"] += 1
        raise

    elapsed_ms = (time.perf_counter() - start) * 1000.0
    _metrics["latency_ms_samples"].append(elapsed_ms)

    runtime = payload.get("runtime", {}) if isinstance(payload, dict) else {}
    generation_mode = str(runtime.get("generation_mode", "unknown"))
    retrieval_mode = str(runtime.get("retrieval_mode", "unknown"))
    _metrics["generation_mode_count"][generation_mode] += 1
    _metrics["retrieval_mode_count"][retrieval_mode] += 1
    if runtime.get("quality_gate_blocked"):
        _metrics["quality_gate_blocked_total"] += 1

    await _save_session_async(session_id, session, user_id=user_id)

    return {"session_id": session_id, **payload}


@app.post("/api/chat/stream")
async def chat_stream(request: ChatRequest, req: Request = None, user_id: int | None = Depends(get_current_user_id)) -> StreamingResponse:
    """SSE streaming endpoint — sends events as they become available."""
    cleaned_prompt = _validate_chat_prompt(request.prompt)
    request_id = getattr(req.state, "request_id", None) if req else None
    session_id = request.session_id or str(uuid4())
    session = _get_or_create_session(session_id, user_id=user_id)
    _metrics["chat_requests_total"] += 1
    start = time.perf_counter()

    async def event_generator():
        try:
            from app.graph import multi_agent_graph
            
            initial_state = {
                "prompt": cleaned_prompt,
                "session": session,
                "request_id": request_id,
                "revision_count": 0
            }
            
            final_state = None
            async for event in multi_agent_graph.astream(initial_state):
                for node_name, node_state in event.items():
                    # Send thinking event for UI
                    yield f"event: thinking\ndata: {json.dumps({'step': node_name}, ensure_ascii=False)}\n\n"
                    final_state = node_state
                    
            if not final_state:
                raise Exception("Graph execution yielded no final state")
                
            payload = {
                "parsed": final_state.get("parsed", {}),
                "retrieved": final_state.get("retrieved_docs", []),
                "generated": {
                    "title": final_state.get("title", ""),
                    "outline": final_state.get("outline", []),
                    "draft": final_state.get("draft", ""),
                    "sources_used": final_state.get("sources_used", [])
                },
                "runtime": {
                    "quality_gate_blocked": bool(final_state.get("editor_feedback")),
                    "generation_mode": "multi-agent",
                    "retrieval_mode": "hybrid"
                }
            }
            
            session.add_turn(
                cleaned_prompt,
                "",  
                parsed_intent=payload["parsed"].get("intent", ""),
                parsed_topic=payload["parsed"].get("topic", ""),
                generated_draft=payload["generated"]["draft"],
            )
            
        except Exception as exc:
            _metrics["chat_errors_total"] += 1
            yield f"event: error\ndata: {json.dumps({'error': str(exc)}, ensure_ascii=False)}\n\n"
            return

        elapsed_ms = (time.perf_counter() - start) * 1000.0
        _metrics["latency_ms_samples"].append(elapsed_ms)

        runtime = payload.get("runtime", {}) if isinstance(payload, dict) else {}
        generation_mode = str(runtime.get("generation_mode", "unknown"))
        retrieval_mode = str(runtime.get("retrieval_mode", "unknown"))
        _metrics["generation_mode_count"][generation_mode] += 1
        _metrics["retrieval_mode_count"][retrieval_mode] += 1
        if runtime.get("quality_gate_blocked"):
            _metrics["quality_gate_blocked_total"] += 1

        # Send metadata event
        meta = {
            "session_id": session_id,
            "parsed": payload.get("parsed"),
            "retrieval_meta": payload.get("retrieved", []),
            "runtime": runtime,
            "latency_ms": round(elapsed_ms, 1),
        }
        yield f"event: meta\ndata: {json.dumps(meta, ensure_ascii=False)}\n\n"

        # Send complete generated content
        result = {"session_id": session_id, **payload}
        yield f"event: done\ndata: {json.dumps(result, ensure_ascii=False)}\n\n"

        # Persist session history
        await _save_session_async(session_id, session, user_id=user_id)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/api/reports")
def create_report(request: SaveReportRequest) -> dict:
    report = save_report(
        prompt=request.prompt,
        title=request.generated.title,
        outline=request.generated.outline,
        draft=request.generated.draft,
        sources_used=request.generated.sources_used,
        sections=[section.model_dump() for section in request.generated.sections],
        session_id=request.session_id,
    )
    return {"report": report}


@app.get("/api/reports")
def reports(limit: int = 50) -> dict:
    return {"reports": list_reports(limit=limit)}


@app.get("/api/reports/{report_id}")
def report_detail(report_id: str) -> dict:
    report = get_report(report_id)
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    return {"report": report}


@app.delete("/api/reports/{report_id}")
def report_delete(report_id: str) -> dict:
    deleted = delete_report(report_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Report not found")
    return {"status": "deleted", "report_id": report_id}


@app.patch("/api/reports/{report_id}/status")
def report_status_update(report_id: str, request: UpdateReportStatusRequest) -> dict:
    report, error = update_report_status(report_id, request.status)
    if error == "not_found":
        raise HTTPException(status_code=404, detail="Report not found")
    if error == "invalid_transition":
        raise HTTPException(
            status_code=409,
            detail="Invalid status transition. Use Draft -> Reviewed -> Approved.",
        )
    if error == "invalid_status":
        raise HTTPException(status_code=400, detail="Invalid report status")
    return {"report": report}


@app.post("/api/reports/{report_id}/publish")
def report_publish(report_id: str, output_format: Literal["markdown", "html"] = "markdown") -> dict:
    report = get_report(report_id)
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    status = str(report.get("status", "Draft"))
    if status != "Approved":
        raise HTTPException(
            status_code=409,
            detail="Only Approved reports can be published.",
        )

    try:
        output = build_publish_output(report, output_format=output_format)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    return {
        "report_id": report_id,
        "status": "published",
        "report_status": status,
        "output": output,
    }


@app.get("/api/source-analytics")
def source_analytics() -> dict:
    try:
        return fetch_source_analytics()
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc


@app.get("/api/knowledge/health")
def knowledge_health() -> dict:
    try:
        return fetch_knowledge_health()
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc


# ── Admin: On-Demand Ingestion API ─────────────────────
_ingestion_state: dict[str, Any] = {
    "running": False,
    "last_result": None,
    "last_run_at": None,
    "last_error": None,
}


def _require_admin_key(x_admin_key: str = Header(default="")) -> None:
    """Verify admin API key. Raises 403 if invalid."""
    expected = settings.admin_api_key.strip()
    if not expected:
        raise HTTPException(status_code=403, detail="Admin API key not configured on server.")
    if x_admin_key != expected:
        raise HTTPException(status_code=403, detail="Invalid admin API key.")


@app.post("/api/admin/ingest")
async def admin_trigger_ingestion(
    x_admin_key: str = Header(default=""),
) -> dict:
    """Trigger a knowledge base re-ingestion job in the background."""
    _require_admin_key(x_admin_key)

    if _ingestion_state["running"]:
        raise HTTPException(status_code=409, detail="Ingestion job is already running.")

    from app.worker import run_ingestion_job

    async def _run():
        _ingestion_state["running"] = True
        _ingestion_state["last_error"] = None
        try:
            result = await asyncio.to_thread(run_ingestion_job)
            _ingestion_state["last_result"] = result
            _ingestion_state["last_run_at"] = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
        except Exception as exc:
            _ingestion_state["last_error"] = str(exc)
        finally:
            _ingestion_state["running"] = False

    asyncio.create_task(_run())
    return {"status": "started", "message": "Ingestion job started in background."}


@app.get("/api/admin/ingest/status")
def admin_ingestion_status(
    x_admin_key: str = Header(default=""),
) -> dict:
    """Get the status of the last ingestion job."""
    _require_admin_key(x_admin_key)
    return {
        "running": _ingestion_state["running"],
        "last_run_at": _ingestion_state["last_run_at"],
        "last_result": _ingestion_state["last_result"],
        "last_error": _ingestion_state["last_error"],
    }
